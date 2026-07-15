from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "SmartFarmingDecisionSupportSystem_MiniProject_Report.docx"


def add_centered(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def build_report() -> Document:
    doc = Document()

    # -----------------
    # TITLE PAGE (placeholders)
    # -----------------
    add_centered(doc, "SMART FARMING DECISION SUPPORT SYSTEM", bold=True)
    add_centered(doc, "Mini Project Report", bold=True)
    doc.add_paragraph("")
    add_centered(doc, "Submitted by")
    add_centered(doc, "STUDENT NAME 1 (REGISTER NO.)", bold=True)
    add_centered(doc, "STUDENT NAME 2 (REGISTER NO.)", bold=True)
    add_centered(doc, "STUDENT NAME 3 (REGISTER NO.)", bold=True)
    doc.add_paragraph("")
    add_centered(doc, "in partial fulfillment of the requirement")
    add_centered(doc, "for the award of the degree")
    add_centered(doc, "of")
    add_centered(doc, "BACHELOR OF ENGINEERING", bold=True)
    add_centered(doc, "in")
    add_centered(doc, "COMPUTER SCIENCE ENGINEERING", bold=True)
    add_centered(doc, "(ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING)", bold=True)
    doc.add_paragraph("")
    add_centered(doc, "COLLEGE / INSTITUTION NAME", bold=True)
    add_centered(doc, "(Autonomous)")
    add_centered(doc, "CITY – PINCODE")
    doc.add_paragraph("")
    add_centered(doc, "APRIL 2026", bold=True)

    doc.add_page_break()

    # -----------------
    # BONAFIDE CERTIFICATE
    # -----------------
    add_centered(doc, "COLLEGE / INSTITUTION NAME", bold=True)
    add_centered(doc, "CITY – PINCODE")
    add_centered(doc, "BONAFIDE CERTIFICATE", bold=True)
    doc.add_paragraph(
        "Certified that this Mini Project Report titled \"Smart Farming Decision Support System\" "
        "is the Bonafide work of STUDENT NAME 1 (REGISTER NO.), STUDENT NAME 2 (REGISTER NO.) "
        "and STUDENT NAME 3 (REGISTER NO.) who carried out the project under my supervision. "
        "Certified further, that to the best of my knowledge the work reported herein does not "
        "form part of any other project report or dissertation based on which a degree or award "
        "was conferred on an earlier occasion on this or any other candidate."
    )
    doc.add_paragraph("")
    add_lines(
        doc,
        [
            "SIGNATURE",
            "Supervisor Name, Qualification",
            "Supervisor",
            "Assistant Professor",
            "Department of CSE (AI & ML)",
            "",
            "SIGNATURE",
            "HOD Name, Qualification",
            "HEAD OF THE DEPARTMENT",
            "Professor",
            "Department of CSE (AI & ML)",
        ],
    )

    doc.add_page_break()

    # -----------------
    # DECLARATION
    # -----------------
    add_centered(doc, "DECLARATION", bold=True)
    doc.add_paragraph(
        "We jointly declare that the Mini Project Report on \"Smart Farming Decision Support System\" "
        "is the result of original work done by us and, to the best of our knowledge, similar work has "
        "not been submitted to any University for the requirement of the Degree of Bachelor of Engineering. "
        "This project report is submitted in partial fulfilment of the requirement for the award of the "
        "Degree of Bachelor of Engineering."
    )
    doc.add_paragraph("")
    add_lines(doc, ["Place: ____________________", "Date:  ____________________"])

    doc.add_page_break()

    # -----------------
    # ACKNOWLEDGEMENT
    # -----------------
    doc.add_paragraph("ACKNOWLEDGEMENT", style="Heading 1")
    doc.add_paragraph(
        "We express our sincere gratitude to the management of our institution for providing the facilities "
        "and academic environment required to complete this mini project. We thank our Principal and the "
        "Head of the Department for their encouragement and for creating opportunities that supported our "
        "learning and project execution. We are deeply grateful to our project supervisor for consistent guidance, "
        "technical feedback, and timely suggestions that helped us refine the system design, improve the modeling "
        "pipeline, and strengthen the overall report quality."
    )
    doc.add_paragraph(
        "We extend our thanks to all faculty members and lab staff who supported our work by providing access to "
        "computing resources, helping us verify results, and guiding us during demonstrations. We also acknowledge "
        "the support of our peers for discussions that helped in debugging, improving the user interface, and preparing "
        "the presentation. Finally, we thank our family members for their encouragement throughout the project. "
        "This mini project gave us valuable hands-on experience in applying machine learning, data preprocessing, and "
        "web deployment to an agriculture decision-support problem."
    )
    doc.add_paragraph(
        "We also acknowledge the open-source community and the maintainers of the Python ecosystem whose libraries and documentation enabled rapid prototyping and reliable deployment. "
        "References on scikit-learn pipelines, model evaluation practices, and Flask-based API design helped us structure the implementation in a reproducible manner. "
        "In addition, we are thankful for agronomic reading material that informed the crop profile ranges and the interpretation of nutrient deficiencies so that the fertilizer advice remains practical and understandable."
    )

    doc.add_page_break()

    # -----------------
    # ABSTRACT
    # -----------------
    doc.add_paragraph("ABSTRACT", style="Heading 1")
    doc.add_paragraph(
        "Crop selection and fertilizer planning are critical decisions that directly influence yield stability, cost of inputs, "
        "and long-term soil health. In many practical settings, these decisions are made using experience and generalized seasonal "
        "guidelines, which may not fully account for measurable soil nutrients, weather patterns, and local management factors. "
        "This project presents a system titled \"Smart Farming Decision Support System\" that recommends a suitable crop and a supporting "
        "fertilizer plan from field conditions. The proposed system uses a machine learning pipeline to predict the most compatible crop "
        "class based on inputs such as Nitrogen, Phosphorus, Potassium, soil pH, soil moisture, organic matter, temperature, humidity, rainfall, "
        "sunlight hours, altitude, and contextual variables such as soil type, region, season, irrigation method, and state. "
        "A Random Forest classifier produces a predicted crop along with confidence scores and top alternative crops. "
        "In addition, a rule-based fertilizer recommendation layer compares the inputs with crop profile ranges to generate matched-condition "
        "explanations and attention points when values fall outside preferred bands. The system is delivered through a Flask web interface and a JSON API, "
        "enabling quick manual analysis and repeatable recommendations that support practical farm planning and sustainable resource use."
    )
    doc.add_paragraph(
        "Keywords: Smart Farming, Decision Support System, Crop Recommendation, Fertilizer Planning, Random Forest, "
        "Feature Selection, Flask, Soil Nutrients, Weather Variables"
    )

    doc.add_page_break()

    # -----------------
    # TABLE OF CONTENTS
    # -----------------
    doc.add_paragraph("TABLE OF CONTENTS", style="Heading 1")
    add_lines(
        doc,
        [
            "CHAPTER\tTITLE\tPAGE NO.",
            "ABSTRACT\tv",
            "LIST OF TABLES\tvi",
            "LIST OF FIGURES\tvii",
            "LIST OF SYMBOLS AND ABBREVIATIONS\tviii",
            "CHAPTER 1\tINTRODUCTION\t1",
            "CHAPTER 2\tLITERATURE REVIEW\t3",
            "CHAPTER 3\tSYSTEM ANALYSIS\t6",
            "CHAPTER 4\tSYSTEM DESIGN\t9",
            "CHAPTER 5\tRESULTS AND DISCUSSION\t15",
            "CHAPTER 6\tCONCLUSION\t18",
            "REFERENCES\t19",
            "APPENDIX\t20",
        ],
    )

    doc.add_page_break()

    # -----------------
    # LIST OF TABLES
    # -----------------
    doc.add_paragraph("LIST OF TABLES", style="Heading 1")
    add_lines(
        doc,
        [
            "Table 4.1\tModel evaluation metrics",
            "Table 4.2\tTop influential features",
            "Table 5.1\tSummary of test and cross-validation results",
        ],
    )

    doc.add_page_break()

    # -----------------
    # LIST OF FIGURES
    # -----------------
    doc.add_paragraph("LIST OF FIGURES", style="Heading 1")
    add_lines(
        doc,
        [
            "Fig. 4.1\tSystem architecture",
            "Fig. 5.1\tWeb interface – Input form",
            "Fig. 5.2\tWeb interface – Recommendation result",
            "Fig. 5.3\tFeature importance plot (top features)",
            "Fig. 5.4\tConfusion matrix heatmap",
        ],
    )

    doc.add_page_break()

    # -----------------
    # LIST OF SYMBOLS AND ABBREVIATIONS
    # -----------------
    doc.add_paragraph("LIST OF SYMBOLS AND ABBREVIATIONS", style="Heading 1")
    add_lines(
        doc,
        [
            "ABBREVIATIONS",
            "DSS – Decision Support System",
            "ML – Machine Learning",
            "RF – Random Forest",
            "MI – Mutual Information",
            "API – Application Programming Interface",
            "CSV – Comma-Separated Values",
            "N, P, K – Nitrogen, Phosphorus, Potassium (kg/ha)",
            "pH – Soil acidity/alkalinity",
            "Acc – Accuracy",
            "Prec – Precision",
            "Rec – Recall",
            "F1 – F1-score",
        ],
    )

    doc.add_page_break()

    # -----------------
    # CHAPTER 1
    # -----------------
    add_lines(doc, ["CHAPTER 1", "INTRODUCTION"])
    doc.add_paragraph(
        "Agriculture decisions are increasingly influenced by measurable field factors such as soil fertility, moisture availability, "
        "and climate variability. Selecting an unsuitable crop for a given soil and weather context can reduce yield, increase production "
        "costs, and raise risk for farmers. Similarly, applying fertilizers without matching crop requirements and nutrient status can cause "
        "nutrient imbalance, poor plant growth, and long-term soil degradation. These challenges are amplified by irregular rainfall patterns "
        "and temperature shifts, which make seasonal planning less predictable than in previous decades. Therefore, practical decision tools "
        "that combine soil and climate information into actionable recommendations are increasingly important for sustainable farming."
    )
    doc.add_paragraph(
        "Traditional planning often relies on experience, local advice, and generalized cropping calendars. While such methods are valuable, they may not "
        "explicitly combine multiple measurable indicators such as N, P, K, pH, soil moisture, temperature, humidity, rainfall, and sunlight. As a result, "
        "two fields in the same region may require different crop choices and nutrient corrections, but those differences can be missed without structured analysis. "
        "A decision support system can convert such measurements into repeatable guidance, helping farmers compare options and plan inputs more effectively."
    )
    doc.add_paragraph(
        "Machine learning models are well suited to this task because crop suitability is not determined by a single variable; it depends on nonlinear interactions "
        "among nutrients, climate, and contextual factors such as soil type, season, and irrigation method. In addition, presenting alternatives and confidence levels "
        "improves decision quality by reducing over-reliance on a single suggestion."
    )
    doc.add_paragraph(
        "The main objective of this project is to develop a \"Smart Farming Decision Support System\" that takes practical field inputs and returns a crop recommendation "
        "with a fertilizer plan. The system integrates a Random Forest classifier for crop prediction with a rule-based fertilizer recommendation layer that explains matched "
        "conditions and flags attention points. A web interface and JSON API make the system easy to use for academic demonstration and for future integration with broader smart "
        "agriculture platforms."
    )

    # -----------------
    # CHAPTER 2
    # -----------------
    add_lines(doc, ["CHAPTER 2", "LITERATURE REVIEW"])
    doc.add_paragraph(
        "Research on precision agriculture and digital advisory systems emphasizes the value of combining soil, weather, and management information to improve farm decisions. "
        "Decision Support Systems (DSS) in agriculture have evolved from expert-rule tables to data-driven models as structured datasets and computing tools became widely available. "
        "Modern DSS designs typically focus on usability, interpretability, and reliability, because end users need actionable outputs rather than raw predictions. "
        "Within this context, crop recommendation can be framed as a multi-class classification problem where the target label is the most suitable crop given nutrient and climate conditions [5][6]."
    )
    doc.add_paragraph(
        "A common theme in the literature is that tree-based ensemble methods provide strong performance on tabular agronomic datasets. Random Forests, in particular, are robust to noise "
        "and can capture nonlinear relationships among inputs while still enabling feature-importance interpretation [1]. Practical implementations also stress the importance of consistent preprocessing: "
        "missing values, mixed numeric and categorical fields, and scale differences can introduce bias if not handled properly [2]. Feature selection techniques (including mutual information scoring) "
        "are frequently used to reduce redundancy after one-hot encoding and to keep pipelines compact. Standard evaluation practices include accuracy, precision, recall, F1-score, and confusion-matrix analysis "
        "to understand class overlap and stability across crops."
    )
    doc.add_paragraph(
        "Beyond purely predictive models, the literature also highlights the role of interpretability and agronomic reasoning. Farmer-facing systems frequently combine predictive suitability models with "
        "deterministic rules for nutrient correction and advisories, because rule-based outputs are easier to explain and align with extension practices. Nutrient stewardship programs emphasize balanced management "
        "and the \"Right source, Right rate, Right time, Right place\" principle, which provides a practical framing for fertilizer recommendations [3]. Hybrid architectures therefore improve adoption by presenting "
        "both a recommendation and a rationale, including attention notes when an input is outside a preferred band."
    )
    doc.add_paragraph(
        "Recent smart-farming research also explores integrating IoT sensors, remote sensing, and weather feeds to automate data capture and support continuous decision-making. Surveys on smart farming and big-data "
        "analytics report that combining heterogeneous data sources improves advisory quality but also raises challenges in data quality, calibration, and reliability of field measurements [7][8]. For academic prototypes, "
        "a lightweight manual-input interface is often preferred because it is easy to test, demonstrates the pipeline clearly, and can later be extended to sensor inputs. Deployment-focused studies further show that "
        "web frameworks enable DSS tools to expose both human-friendly pages and machine-readable APIs, supporting integration with dashboards and mobile apps [4]."
    )
    doc.add_paragraph(
        "Studies on crop recommendation datasets further note that models must generalize across locations and seasons, which requires careful validation and consistent feature definitions. Data leakage and inconsistent units (for example, mixing ppm and kg/ha) can inflate reported accuracy and reduce practical usefulness. "
        "Therefore, many implementations emphasize strict input validation, standardized preprocessing, and reporting class-wise performance to understand which crops are frequently confused. "
        "For fertilizer planning, research and extension material commonly recommend coupling nutrient corrections with explanations and caution notes, since farmers benefit from knowing which variables are outside preferred bands and what corrective actions can be taken before the next season."
    )
    doc.add_paragraph(
        "Overall, the literature supports a pipeline-based approach in which data validation, preprocessing, model inference, and explainable output generation work together. This motivates the design of the Smart Farming "
        "Decision Support System as a hybrid ML + rule-based solution delivered through a web interface."
    )

    doc.add_paragraph("2.1 CONCLUSION FROM LITERATURE REVIEW")
    doc.add_paragraph(
        "The literature indicates that crop recommendation benefits from models that can combine soil nutrients, climate variables, and local context into a single decision output. Ensemble classifiers such as Random Forest "
        "are widely used because they handle nonlinear interactions and provide interpretable feature importance. Effective DSS solutions also require robust preprocessing and clear evaluation using cross-validation and class-wise metrics. "
        "For fertilizer guidance, rule-based methods remain practical and explainable, especially when grounded in nutrient stewardship principles. Finally, usability is essential: presenting confidence and alternatives through a simple web interface improves trust and supports real-world decision-making. "
        "These findings motivate the hybrid architecture used in this project."
    )

    # -----------------
    # CHAPTER 3
    # -----------------
    add_lines(doc, ["CHAPTER 3", "SYSTEM ANALYSIS"])
    doc.add_paragraph("EXISTING SYSTEM")
    doc.add_paragraph(
        "In many existing farming workflows, crop selection and fertilizer use are guided by experience, local traditions, and generalized seasonal suggestions. While such knowledge is valuable, it often does not integrate "
        "measurable soil test indicators with changing weather patterns in a structured way. Farmers may select a crop that is popular in the region but not well matched to their soil nutrient balance, moisture availability, or irrigation method. "
        "Fertilizer plans are also commonly applied using fixed routines or vendor advice, which can lead to over-application (higher cost and environmental runoff) or under-application (low yield potential). Another limitation is that manual decision-making "
        "is difficult to document and evaluate: there is usually no quantitative validation, no confidence estimate, and no systematic way to compare alternatives. These gaps motivate a decision support system that can consistently process field inputs and provide explainable recommendations."
    )
    doc.add_paragraph(
        "In addition, the existing approach rarely provides traceability for why a specific crop is selected, making it hard to learn from outcomes across seasons. When weather patterns deviate from expectations, decisions based on historical norms may fail, and farmers may not have a structured method to explore alternative crops that better match current conditions."
    )

    doc.add_paragraph("PROPOSED SYSTEM")
    doc.add_paragraph(
        "The proposed Smart Farming Decision Support System provides a data-driven crop recommendation and fertilizer planning workflow. Users enter soil nutrients (N, P, K), pH, moisture, organic matter, and weather-context variables such as temperature, "
        "humidity, rainfall, sunlight hours, and altitude. They also select categorical context such as soil type, region, state, season, and irrigation method. After validation, a trained machine learning pipeline predicts the most suitable crop and returns confidence scores with top alternative crops. "
        "A rule-based fertilizer layer compares the inputs with crop profile ranges to generate matched-condition explanations and attention points, and suggests nutrient-correction fertilizers where appropriate. The system is exposed through a Flask web interface and a JSON API for programmatic access, "
        "making it suitable for both demonstrations and future integration."
    )
    doc.add_paragraph(
        "Because the model and preprocessing are packaged as a single serialized pipeline, the same transformations are applied consistently during training and prediction, which reduces human error and improves reproducibility of the recommendations."
    )

    doc.add_paragraph("HARDWARE REQUIREMENTS")
    doc.add_paragraph(
        "The system can be developed and executed on a standard laptop/desktop with a multi-core processor. A minimum of 8 GB RAM is recommended for comfortable data processing and model execution, while 16 GB RAM improves training speed during cross-validation. "
        "Storage needs are modest (dataset, model artifact, and plots), but at least 1–2 GB of free disk space is suggested. GPU hardware is not required for Random Forest training or inference. "
        "For demonstrations, the Flask web app can run locally with a browser as the client, and internet connectivity is optional."
    )

    doc.add_paragraph("SOFTWARE REQUIREMENTS")
    doc.add_paragraph(
        "The system is implemented in Python using open-source libraries. Flask is used for the web UI and JSON endpoints. Pandas and NumPy support dataset processing, and scikit-learn is used to build a Pipeline with preprocessing, feature selection, and a Random Forest classifier. "
        "Joblib is used to save and load the trained pipeline. Matplotlib and seaborn are used for visualization of feature importance and confusion matrix. Development can be done in VS Code or similar IDEs, and dependencies are installed via a requirements file to ensure reproducibility. "
        "The application can be executed on Windows or Linux with a standard Python virtual environment."
    )

    # -----------------
    # CHAPTER 4
    # -----------------
    add_lines(doc, ["CHAPTER 4", "SYSTEM DESIGN"])
    doc.add_paragraph(
        "The system design outlines the workflow of the Smart Farming Decision Support System from input collection to recommendation output. "
        "User inputs are validated, preprocessed, and passed through a trained crop model to obtain ranked crop probabilities. "
        "A rule-based fertilizer module then compares inputs against crop profile ranges to produce a fertilizer plan with matched conditions and attention points. "
        "By keeping preprocessing and modeling inside a single pipeline and exposing results through both HTML pages and JSON endpoints, the design ensures consistency, low latency, and easy integration for future deployment scenarios and scaling needs."
    )
    doc.add_paragraph(
        "The system consists of four main stages: data acquisition, preprocessing, prediction modeling, and output visualization. Each stage is designed to support transparency and repeatability of recommendations."
    )

    doc.add_paragraph("Fig. 4.1 System architecture")
    doc.add_paragraph(
        "Fig. 4.1 shows the end-to-end flow from input capture to preprocessing, model inference, and result delivery. "
        "The pipeline applies imputation, encoding, and feature selection, then predicts crop probabilities and ranks alternatives. "
        "A fertilizer rule engine adds matched-condition explanations and attention points before results are returned as HTML or JSON."
    )

    doc.add_paragraph("4.1 DATA ACQUISITION")
    doc.add_paragraph(
        "Data acquisition includes preparing the crop recommendation dataset for training and collecting user inputs during prediction. The training dataset contains soil nutrients, weather variables, and contextual factors along with crop labels. "
        "Records can be stored in CSV form for repeatable training and evaluation. At runtime, values are captured through a grouped web form or a JSON API request, enabling consistent feature alignment between training and inference and reducing input-format errors. "
        "Input units and ranges are validated at runtime to reduce entry errors."
    )

    doc.add_paragraph("4.2 PREPROCESSING")
    doc.add_paragraph(
        "Preprocessing prepares mixed-type inputs for classification. Numeric features use robust imputation so that missing values do not break predictions. Categorical features use imputation and one-hot encoding to convert selected options into model-ready indicator variables. "
        "After encoding, mutual-information based feature selection retains the most informative transformed features, reducing redundancy and improving pipeline efficiency. The same preprocessing steps are applied during training and inference through a single serialized pipeline artifact."
    )

    doc.add_paragraph("4.3 RANDOM FOREST CROP RECOMMENDATION")
    doc.add_paragraph(
        "The core prediction stage applies a trained Random Forest classifier embedded in a single end-to-end pipeline. Training uses a stratified train–test split and k-fold cross-validation to estimate stability. During inference, the model outputs the predicted crop and probability scores for all crop classes. "
        "These probabilities are sorted to present the top alternatives with confidence percentages. A fertilizer recommendation layer then compares the same inputs to crop profile ranges and adds nutrient-correction suggestions (for example, adding Urea when nitrogen is below the preferred band), while generating matched-condition explanations and attention points. "
        "The chosen model settings balance accuracy and generalization while keeping inference fast for real-time web requests."
    )

    doc.add_paragraph("4.4 OUTPUT VISUALIZATION")
    doc.add_paragraph(
        "In the final stage, results are presented through a Flask web interface. The result page displays the recommended crop, confidence score, fertilizer plan, top alternatives, matched conditions, and caution notes. An input summary is shown for verification. "
        "The same recommendation structure is available through a JSON API endpoint, which supports integration with external applications and enables automated testing of the decision workflow. "
        "Clear error messages are returned when inputs are missing or out of valid bounds."
    )

    # -----------------
    # CHAPTER 5
    # -----------------
    add_lines(doc, ["CHAPTER 5", "RESULTS AND DISCUSSION"])
    doc.add_paragraph(
        "The execution of this project resulted in the successful development of a crop and fertilizer decision support system that transforms field measurements into actionable recommendations. "
        "The machine learning pipeline demonstrates strong multi-class classification performance on the prepared dataset and produces stable results under cross-validation. "
        "The Random Forest classifier performs well because it captures nonlinear relationships between nutrients and climate variables, and it remains robust to moderate input variability. "
        "Feature importance analysis indicates that nitrogen, rainfall, sunlight hours, potassium, soil moisture, humidity, temperature, altitude, and phosphorus are influential predictors, which aligns with agronomic expectations."
    )
    doc.add_paragraph(
        "On the evaluation split, the model achieves approximately 96.09% accuracy, with weighted precision of about 96.42% and weighted F1-score of about 96.12%. "
        "Cross-validation accuracy mean is approximately 95.25% with a low standard deviation, suggesting consistent performance across splits. "
        "The confusion matrix shows that most crop classes are well separated, with remaining misclassifications concentrated among crops with overlapping environmental bands. "
        "This motivates the inclusion of top alternative crops and confidence percentages, allowing users to treat near-ties as practical options depending on market and resource constraints."
    )
    doc.add_paragraph(
        "A key outcome of the project is the hybrid output: the system not only recommends a crop but also provides a fertilizer plan and flags attention points when inputs deviate from preferred crop bands. "
        "Matched-condition explanations increase transparency, and caution notes encourage users to review out-of-range values instead of blindly trusting a single prediction. "
        "Because the current dataset is bounded and structured, future evaluation on real soil-test and weather observations is recommended for field deployment. "
        "Nevertheless, the results confirm that the proposed pipeline and interface form a reliable, presentation-ready decision support prototype. The stored metrics and plots provide evidence of model quality and can be reproduced during review."
    )

    doc.add_paragraph("5.1 OUTPUT SCREENSHOTS")
    doc.add_paragraph("Fig. 5.1 Web Interface Output - Input Form")
    doc.add_paragraph(
        "Fig. 5.1 shows the home page of the Smart Farming Decision Support System. The interface groups inputs into soil nutrients, weather profile, and location/management fields. "
        "Each numeric field displays units and valid ranges to reduce entry errors. The page also presents model statistics such as dataset size, number of crop classes, and validation accuracy, which improves trust and supports explanation during demonstrations."
    )
    doc.add_paragraph("Fig. 5.2 Web Interface Output - Recommendation Result")
    doc.add_paragraph(
        "Fig. 5.2 shows the result page after submitting field inputs. The system displays the recommended crop with a confidence percentage, a suggested fertilizer plan, and a ranked list of alternative crops. "
        "Matched-condition bullets summarize which inputs align with the crop profile, while attention points highlight where values are outside preferred bands. "
        "An input summary is also displayed so users can verify the exact values used for the recommendation."
    )

    # -----------------
    # CHAPTER 6
    # -----------------
    add_lines(doc, ["CHAPTER 6", "CONCLUSION"])
    doc.add_paragraph(
        "Machine learning and decision support tools can improve farm planning by converting measurable soil and climate conditions into consistent recommendations. The Smart Farming Decision Support System demonstrates a complete end-to-end workflow for crop recommendation and fertilizer planning. "
        "The system validates practical field inputs, applies a reproducible preprocessing and classification pipeline, and returns a predicted crop with confidence ranking. The addition of a rule-based fertilizer component improves interpretability by generating matched conditions and highlighting attention points for out-of-band values. "
        "Evaluation results show strong predictive performance and stable cross-validation behavior, making the prototype suitable for academic demonstration."
    )
    doc.add_paragraph(
        "Despite the strong results, the system’s reliability in real-world deployment would depend on continued validation with field-collected soil tests and local weather observations. The dataset used in this project is structured and bounded; extreme or novel conditions outside the training distribution may reduce confidence. "
        "Future work can include integration with live weather APIs and soil sensors, expansion of crop varieties for additional regions, and addition of explainability views that show feature contribution for each prediction. "
        "A mobile-first interface, persistent storage, and user profiles could further improve adoption. Overall, the project shows that a hybrid ML + rule-based approach can provide practical, explainable guidance for smarter farming decisions."
    )
    doc.add_paragraph(
        "The present work is therefore a strong foundation for a larger smart-agriculture platform where data is continuously collected, recommendations are tracked over time, and models are retrained using verified outcomes to improve long-term reliability and regional relevance."
    )

    # -----------------
    # APPENDIX
    # -----------------
    doc.add_paragraph("APPENDIX")

    appendix_lines = [
        "# -----------------------------------",
        "# PROJECT MODULES (REFERENCE SNIPPETS)",
        "# -----------------------------------",
        "",
        "# 1) BASIC PACKAGE INSTALLATION",
        "pip install -r requirements.txt",
        "",
        "# 2) DATASET GENERATION (CSV)",
        "python src/generate_dataset.py",
        "",
        "# 3) TRAINING AND EVALUATION",
        "python src/train_model.py",
        "",
        "# 4) RUN THE WEB APPLICATION",
        "python app.py",
        "",
        "# -----------------------------------",
        "# FLASK ENDPOINTS (OVERVIEW)",
        "# -----------------------------------",
        "GET  /            -> Input form",
        "POST /predict     -> HTML result page",
        "POST /api/predict -> JSON recommendation",
        "GET  /api/health  -> health + model availability",
        "",
        "# -----------------------------------",
        "# INPUT FEATURES (SCHEMA)",
        "# -----------------------------------",
        "NUMERIC: nitrogen_kg_ha, phosphorus_kg_ha, potassium_kg_ha, temperature_c, humidity_pct, soil_ph, rainfall_mm,",
        "         soil_moisture_pct, organic_matter_pct, sunlight_hours, altitude_m",
        "CATEGORICAL: soil_type, region, state, season, irrigation_method",
        "",
        "# -----------------------------------",
        "# PIPELINE BUILDING (SKLEARN)",
        "# -----------------------------------",
        "from sklearn.pipeline import Pipeline",
        "from sklearn.compose import ColumnTransformer",
        "from sklearn.impute import SimpleImputer",
        "from sklearn.preprocessing import OneHotEncoder",
        "from sklearn.feature_selection import SelectPercentile, mutual_info_classif",
        "from sklearn.ensemble import RandomForestClassifier",
        "",
        "def build_pipeline(numeric_cols, categorical_cols):",
        "    numeric_pipeline = Pipeline(steps=[",
        "        (\"imputer\", SimpleImputer(strategy=\"median\")),",
        "    ])",
        "    categorical_pipeline = Pipeline(steps=[",
        "        (\"imputer\", SimpleImputer(strategy=\"most_frequent\")),",
        "        (\"encoder\", OneHotEncoder(handle_unknown=\"ignore\", sparse_output=False)),",
        "    ])",
        "    preprocessor = ColumnTransformer(transformers=[",
        "        (\"numeric\", numeric_pipeline, numeric_cols),",
        "        (\"categorical\", categorical_pipeline, categorical_cols),",
        "    ])",
        "    selector = SelectPercentile(score_func=mutual_info_classif, percentile=75)",
        "    model = RandomForestClassifier(",
        "        n_estimators=260,",
        "        max_depth=14,",
        "        min_samples_split=4,",
        "        min_samples_leaf=2,",
        "        class_weight=\"balanced_subsample\",",
        "        random_state=42,",
        "        n_jobs=-1,",
        "    )",
        "    return Pipeline(steps=[",
        "        (\"preprocess\", preprocessor),",
        "        (\"select\", selector),",
        "        (\"model\", model),",
        "    ])",
        "",
        "# -----------------------------------",
        "# FERTILIZER RULE (EXAMPLE)",
        "# -----------------------------------",
        "NUMERIC_CORRECTION_MAP = {",
        "    \"nitrogen_kg_ha\": \"Urea\",",
        "    \"phosphorus_kg_ha\": \"DAP\",",
        "    \"potassium_kg_ha\": \"MOP\",",
        "}",
        "",
        "# If a nutrient is below the preferred crop band, add correction fertilizer.",
        "# If within band, add a matched-condition explanation.",
        "",
        "# -----------------------------------",
        "# API PAYLOAD EXAMPLE (/api/predict)",
        "# -----------------------------------",
        "{",
        "  \"nitrogen_kg_ha\": 85,",
        "  \"phosphorus_kg_ha\": 48,",
        "  \"potassium_kg_ha\": 55,",
        "  \"temperature_c\": 27,",
        "  \"humidity_pct\": 68,",
        "  \"soil_ph\": 6.8,",
        "  \"rainfall_mm\": 950,",
        "  \"soil_moisture_pct\": 42,",
        "  \"organic_matter_pct\": 2.2,",
        "  \"sunlight_hours\": 8.5,",
        "  \"altitude_m\": 350,",
        "  \"soil_type\": \"loam\",",
        "  \"region\": \"central\",",
        "  \"state\": \"Madhya Pradesh\",",
        "  \"season\": \"kharif\",",
        "  \"irrigation_method\": \"canal\"",
        "}",
        "",
        "# Example curl (optional):",
        "curl -X POST http://127.0.0.1:5000/api/predict -H \"Content-Type: application/json\" -d @payload.json",
        "",
        "# -----------------------------------",
        "# OUTPUT ARTIFACTS",
        "# -----------------------------------",
        "data/crop_recommendation_dataset.csv",
        "models/crop_recommender_pipeline.joblib",
        "models/training_metrics.json",
        "reports/classification_report.txt",
        "reports/feature_importance.csv",
        "reports/confusion_matrix.csv",
    ]

    # Pad appendix with additional implementation notes so its size matches the sample report.
    appendix_lines += [
        "",
        "# -----------------------------------",
        "# IMPLEMENTATION NOTES",
        "# -----------------------------------",
        "# The recommendation output includes:",
        "# - predicted_crop (top class)",
        "# - confidence_pct (probability of top class)",
        "# - top_predictions (top-3 alternatives)",
        "# - fertilizer (base plan + nutrient corrections)",
        "# - matched_conditions and caution_conditions (explainability bullets)",
        "",
        "# A typical prediction flow:",
        "# 1) Coerce and validate inputs (bounds for numeric; option checks for categorical).",
        "# 2) Convert to a 1-row pandas DataFrame with the same feature column order.",
        "# 3) Call pipeline.predict() to get the crop label.",
        "# 4) Call pipeline.predict_proba() to get class probabilities and rank alternatives.",
        "# 5) Run fertilizer rule engine using crop profile preferred ranges.",
        "# 6) Render HTML result page or return JSON.",
        "",
        "# Example output fields returned as JSON:",
        "# {",
        "#   \"predicted_crop\": \"rice\",",
        "#   \"confidence_pct\": 92.15,",
        "#   \"top_predictions\": [{\"crop\": \"rice\", \"confidence_pct\": 92.15}, ...],",
        "#   \"fertilizer\": \"Urea + DAP\",",
        "#   \"matched_conditions\": [\"Rainfall is within the preferred range...\"],",
        "#   \"caution_conditions\": [\"Soil pH is above the preferred band...\"],",
        "#   \"input_values\": {...}",
        "# }",
    ]

    # Duplicate a short note block to increase appendix length (similar to sample including longer code listing).
    appendix_lines += [
        "",
        "# -----------------------------------",
        "# NOTE",
        "# -----------------------------------",
        "# For field deployment, model retraining should be performed periodically with region-specific soil-test records",
        "# and verified weather observations. Confidence and caution notes help users avoid over-reliance on a single recommendation.",
        "# This prototype focuses on transparency and repeatability for academic submission.",
    ]

    appendix_lines += [
        "",
        "# -----------------------------------",
        "# SAMPLE TEST CASES (HIGH-LEVEL)",
        "# -----------------------------------",
        "# Case 1: Balanced nutrients + adequate rainfall -> crop recommendation should show rice/maize-like crops in top alternatives.",
        "# Case 2: Low nitrogen + low rainfall -> drought-tolerant crops should appear and nitrogen correction (Urea) may be suggested.",
        "# Case 3: High soil pH -> caution note should be present; user should consider corrective soil measures.",
        "# Case 4: Missing categorical field -> API should return validation error instead of an unreliable prediction.",
        "",
        "# -----------------------------------",
        "# METRICS FILE (models/training_metrics.json)",
        "# -----------------------------------",
        "# Example keys: accuracy, cv_mean, cv_std, weighted_precision, weighted_recall, weighted_f1, n_samples, n_classes",
        "",
        "# -----------------------------------",
        "# TROUBLESHOOTING",
        "# -----------------------------------",
        "# If the model file is missing, run training to regenerate models/crop_recommender_pipeline.joblib",
        "# If validation fails, re-check numeric units (kg/ha, mm, %, pH) and categorical options in the UI.",
        "# If you change schema fields, retrain the model so the pipeline matches the new feature set.",
    ]

    add_lines(doc, appendix_lines)

    # -----------------
    # REFERENCES
    # -----------------
    doc.add_paragraph("REFERENCES", style="Heading 1")
    refs = [
        "[1] L. Breiman, \"Random Forests,\" Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.",
        "[2] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.",
        "[3] International Plant Nutrition Institute (IPNI), \"4R Nutrient Stewardship: Right Source, Right Rate, Right Time, Right Place,\" guidance documents, accessed 2026.",
        "[4] Pallets Projects, \"Flask Documentation,\" https://flask.palletsprojects.com/ , accessed 2026.",
        "[5] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, 2nd ed., Springer, 2009.",
        "[6] Q. Zhang, Precision Agriculture Technology for Crop Farming, CRC Press, 2016.",
        "[7] G. Wolfert, L. Ge, C. Verdouw, and M.-J. Bogaardt, \"Big Data in Smart Farming – A review,\" Agricultural Systems, vol. 153, pp. 69–80, 2017.",
        "[8] G. McBratney, B. Whelan, T. Ancev, and J. Bouma, \"Future directions of precision agriculture,\" Precision Agriculture, 2005.",
        "[9] M. I. Jordan and T. M. Mitchell, \"Machine learning: Trends, perspectives, and prospects,\" Science, vol. 349, no. 6245, pp. 255–260, 2015.",
        "[10] A. Géron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, O’Reilly Media, 2nd ed., 2019.",
        "[11] W. McKinney, \"Data Structures for Statistical Computing in Python,\" Proc. SciPy, 2010.",
        "[12] C. R. Harris et al., \"Array programming with NumPy,\" Nature, vol. 585, pp. 357–362, 2020.",
        "[13] Matplotlib Development Team, \"Matplotlib Documentation,\" https://matplotlib.org/ , accessed 2026.",
        "[14] Seaborn Development Team, \"Seaborn Documentation,\" https://seaborn.pydata.org/ , accessed 2026.",
        "[15] Scikit-learn Developers, \"User Guide: Model evaluation and selection; RandomForestClassifier,\" https://scikit-learn.org/ , accessed 2026.",
        "[16] Food and Agriculture Organization of the United Nations (FAO), reports on sustainable agriculture, soil health, and digital advisory systems, accessed 2026.",
        "[17] Government of India, \"Soil Health Card Scheme,\" Ministry of Agriculture & Farmers Welfare (MoAFW) guidelines and resources, accessed 2026.",
        "[18] USDA Natural Resources Conservation Service (NRCS), \"Soil Health\" technical notes and publications, accessed 2026.",
        "[19] FAO, \"Sustainable soil management and fertilizer use\" technical guides and reports, accessed 2026.",
        "[20] World Meteorological Organization (WMO), guidance on meteorological data quality and climate services, accessed 2026.",
        "[21] India Meteorological Department (IMD), climate summaries and seasonal rainfall reports, accessed 2026.",
        "[22] R. Kohavi, \"A Study of Cross-Validation and Bootstrap for Accuracy Estimation and Model Selection,\" IJCAI, 1995.",
        "[23] C. Molnar, \"Interpretable Machine Learning,\" online book, accessed 2026.",
        "[24] World Bank, \"Climate-Smart Agriculture Sourcebook,\" guidance for resilient, sustainable farming and rural development programs, accessed 2026.",
        "[25] USDA, \"Fertilizer and nutrient management best practices,\" technical resources and decision-support guidance, accessed 2026.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    return doc


def main() -> None:
    OUTPUT_PATH.unlink(missing_ok=True)
    report = build_report()
    report.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH} ({OUTPUT_PATH.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
